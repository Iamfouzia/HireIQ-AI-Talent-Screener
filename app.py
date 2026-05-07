import os
import io
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from flask import Flask, render_template, request, jsonify, session, redirect, url_for
from werkzeug.utils import secure_filename
from rag_pipeline import analyze_resume, extract_text
from graph_builder import build_knowledge_graph
from dotenv import load_dotenv
from concurrent.futures import ThreadPoolExecutor, as_completed
from fpdf import FPDF
import openpyxl
from docx import Document

load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY', 'hireiq_secret_2024')
app.config['UPLOAD_FOLDER']      = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 200 * 1024 * 1024

ALLOWED_EXTENSIONS = {'pdf'}
os.makedirs('uploads', exist_ok=True)


def allowed_file(filename):
    """Return True only for PDF files."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def login_required(f):
    """Decorator to protect routes — redirect to login if not logged in."""
    from functools import wraps
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('logged_in'):
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated


# ROUTE 1: Login page
@app.route('/login', methods=['GET', 'POST'])
def login():
    """Show login form and handle login submission."""
    error = None
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '').strip()

        # Check credentials against .env values
        if username == os.getenv('HR_USERNAME') and password == os.getenv('HR_PASSWORD'):
            session['logged_in'] = True
            session['username']  = username
            return redirect(url_for('index'))
        else:
            error = 'Invalid username or password'

    return render_template('login.html', error=error)


# ROUTE 2: Logout
@app.route('/logout')
def logout():
    """Clear session and redirect to login."""
    session.clear()
    return redirect(url_for('login'))


# ROUTE 3: Main page — protected
@app.route('/')
@login_required
def index():
    return render_template('index.html', username=session.get('username'))


# ROUTE 4: Analyze resumes — protected
@app.route('/analyze', methods=['POST'])
@login_required
def analyze():
    """Accept PDF resumes + job description. Process and return ranked results."""

    job_description = request.form.get('job_description', '').strip()
    if not job_description:
        return jsonify({"success": False, "error": "Job description is required"}), 400

    files = request.files.getlist('resumes')
    if not files:
        return jsonify({"success": False, "error": "No files uploaded"}), 400

    # Save valid PDFs to disk
    saved_files = []
    for file in files:
        if file and allowed_file(file.filename):
            filename       = secure_filename(file.filename)
            candidate_name = filename.replace('.pdf', '')
            filepath       = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            saved_files.append((filepath, candidate_name))

    if not saved_files:
        return jsonify({"success": False, "error": "No valid PDF files found"}), 400

    def process_one(filepath, candidate_name):
        """Run RAG + knowledge graph for a single resume."""
        try:
            result     = analyze_resume(filepath, job_description, candidate_name)
            text       = extract_text(filepath)
            graph_data = build_knowledge_graph(text, candidate_name)
            result['graph'] = graph_data
            return result
        except Exception as e:
            return {
                "name": candidate_name,
                "score": 0,
                "analysis": {"Error": str(e)},
                "graph": {"nodes": [], "edges": []}
            }

    # Process resumes one at a time to avoid FAISS conflicts
    results = []
    with ThreadPoolExecutor(max_workers=1) as executor:
        futures = {
            executor.submit(process_one, fp, cn): cn
            for fp, cn in saved_files
        }
        for future in as_completed(futures):
            results.append(future.result())

    results.sort(key=lambda x: x.get('score', 0), reverse=True)

    return jsonify({"success": True, "results": results, "total": len(results)})


# ROUTE 5: Export as PDF — protected
@app.route('/export/pdf', methods=['POST'])
@login_required
def export_pdf():
    """Generate downloadable PDF report."""
    data     = request.get_json()
    results  = data.get('results', [])
    job_desc = data.get('job_description', 'N/A')

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", 'B', 16)
    pdf.cell(0, 10, "HireIQ — Resume Screening Results", ln=True, align='C')
    pdf.ln(4)
    pdf.set_font("Helvetica", size=10)
    pdf.cell(0, 8, f"Role: {job_desc[:80]}", ln=True)
    pdf.ln(4)

    pdf.set_font("Helvetica", 'B', 11)
    pdf.set_fill_color(20, 20, 20)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(15, 10, "Rank",      border=1, fill=True)
    pdf.cell(80, 10, "Candidate", border=1, fill=True)
    pdf.cell(35, 10, "Score",     border=1, fill=True)
    pdf.cell(60, 10, "Status",    border=1, fill=True, ln=True)

    pdf.set_font("Helvetica", size=10)
    pdf.set_text_color(0, 0, 0)
    for i, r in enumerate(results):
        score  = r.get('score', 0)
        status = "Shortlisted" if score >= 70 else "Review" if score >= 50 else "Rejected"
        pdf.set_fill_color(240, 240, 240) if i % 2 == 0 else pdf.set_fill_color(255, 255, 255)
        pdf.cell(15, 9, str(i + 1),             border=1, fill=True)
        pdf.cell(80, 9, r.get('name', '')[:35], border=1, fill=True)
        pdf.cell(35, 9, f"{score}%",            border=1, fill=True)
        pdf.cell(60, 9, status,                 border=1, fill=True, ln=True)

    pdf_bytes = bytes(pdf.output())
    return app.response_class(
        pdf_bytes,
        mimetype='application/pdf',
        headers={"Content-Disposition": "attachment; filename=hireiq_results.pdf"}
    )


# ROUTE 6: Export as Excel — protected
@app.route('/export/excel', methods=['POST'])
@login_required
def export_excel():
    """Generate downloadable Excel report."""
    data     = request.get_json()
    results  = data.get('results', [])
    job_desc = data.get('job_description', 'N/A')

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "HireIQ Results"
    ws.append(["Rank", "Candidate Name", "Match Score (%)", "Status", "Job Description"])

    for i, r in enumerate(results):
        score  = r.get('score', 0)
        status = "Shortlisted" if score >= 70 else "Review" if score >= 50 else "Rejected"
        ws.append([i + 1, r.get('name', ''), score, status, job_desc if i == 0 else ""])

    ws.column_dimensions['A'].width = 8
    ws.column_dimensions['B'].width = 30
    ws.column_dimensions['C'].width = 18
    ws.column_dimensions['D'].width = 15
    ws.column_dimensions['E'].width = 50

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return app.response_class(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={"Content-Disposition": "attachment; filename=hireiq_results.xlsx"}
    )


# ROUTE 7: Export as Word — protected
@app.route('/export/word', methods=['POST'])
@login_required
def export_word():
    """Generate downloadable Word report."""
    data     = request.get_json()
    results  = data.get('results', [])
    job_desc = data.get('job_description', 'N/A')

    doc = Document()
    doc.add_heading('HireIQ — Resume Screening Results', 0)
    doc.add_paragraph(f"Job Role: {job_desc}")
    doc.add_paragraph("")

    table          = doc.add_table(rows=1, cols=4)
    table.style    = 'Table Grid'
    header         = table.rows[0].cells
    header[0].text = 'Rank'
    header[1].text = 'Candidate'
    header[2].text = 'Score'
    header[3].text = 'Status'

    for i, r in enumerate(results):
        score  = r.get('score', 0)
        status = "Shortlisted" if score >= 70 else "Review" if score >= 50 else "Rejected"
        row    = table.add_row().cells
        row[0].text = str(i + 1)
        row[1].text = r.get('name', '')
        row[2].text = f"{score}%"
        row[3].text = status

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return app.response_class(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={"Content-Disposition": "attachment; filename=hireiq_results.docx"}
    )


# ROUTE 8: Send email — protected
@app.route('/send-email', methods=['POST'])
@login_required
def send_email():
    """Send shortlist email to HR."""
    data       = request.get_json()
    hr_email   = data.get('hr_email', '').strip()
    candidates = data.get('candidates', [])
    job_desc   = data.get('job_description', '')

    if not hr_email:
        return jsonify({"success": False, "error": "HR email is required"}), 400
    if not candidates:
        return jsonify({"success": False, "error": "No candidates selected"}), 400

    body  = f"Job Role: {job_desc}\n\n"
    body += "Shortlisted Candidates:\n" + "-" * 40 + "\n"
    for i, c in enumerate(candidates):
        body += f"{i+1}. {c['name']} — {c['score']}% Match\n"
    body += "-" * 40 + "\nGenerated by HireIQ — AI Powered Talent Screener"

    msg            = MIMEMultipart()
    msg['From']    = os.getenv('EMAIL_ADDRESS')
    msg['To']      = hr_email
    msg['Subject'] = f"HireIQ Shortlist — {job_desc[:40]}"
    msg.attach(MIMEText(body, 'plain'))


    try:
        import resend
        resend.api_key = os.getenv('RESEND_API_KEY')
        resend.Emails.send({
            "from": "onboarding@resend.dev",
            "to": hr_email,
            "subject": f"HireIQ Shortlist — {job_desc[:40]}",
            "text": body
        })
        return jsonify({"success": True, "message": "Email sent successfully!"})
    except Exception as e:
        print(f"[Email] Error: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))